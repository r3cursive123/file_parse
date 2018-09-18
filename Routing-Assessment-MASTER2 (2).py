from libnmap.process import NmapProcess   #Import of necessary modules
#from time import sleep
from libnmap.parser import NmapParser
from docx import Document

print 'Please enter the NAME (ex. Production1) of the source network you are validating from:'
sourceNetworkName =  raw_input()
print 'Please enter the appropite network ID (ex. 192.168.1.0/24)'
sourceNetwork =  raw_input()
print '..................'

"""destinations = {"staging":"192.168.159.0/24","production1":"127.0.0.1", "extra":"127.0.0.1"}"""

destinations = {"staging":"192.168.159.0/24","production1":"127.0.0.1", "extra":"127.0.0.1",}

scanOptions = "-sS -O -v -p T:1-99,121-1025,1068,1126,1433,1434,1720,1764,2049,2380-2390,3268,3269,3389,5357,7000-7200,7220-7796,7800-9010"

outPutFile = open(sourceNetworkName + " " + sourceNetwork.replace('/','-'), "w")  # This is the output file defined

document = Document ()

allReadyPrinted = False

for element in destinations.keys():   #  This is prohibiting the scan and removing the destination network if the destination network matches the source network
	if destinations[element] == sourceNetwork:
		del(destinations[element])


for dest in destinations.keys():   #  This is the loop to process the keys (items) in the dictionary
	
	
	#for NmapTask in NmapProcess.tasks:
    #print scanned_hosts
	#print dir(NmapProcess.tasks)
	#print dir (NmapProcess.tasks)
    #print (NmapProcess.targets)


	nm = NmapProcess(destinations[dest], options=scanOptions)    # This is the running of the scan on the defined destinations in the defenition
	rc = nm.run()
	if nm.rc == 0:
	    #result = nm.stdout
	    #print nm.stdout
		
		nmap_report = NmapParser.parse(nm.stdout)  # This is the parseing of the results stdout.


		for scanned_hosts in nmap_report.hosts:  # This is a loop for host identified in the report
	    
	    	#print dir (scanned_hosts)
		#for open_ports in nmap_report.
			#print scanned_hosts.ipv4
			#print scanned_hosts.get_open_ports()

			output = scanned_hosts.ipv4 	# This is the results of the targets ipv4 address
			#statusOutput = scanned_hosts.
			for port in scanned_hosts.get_open_ports():      # This calls out the open ports in the scan
				output += " " + "PORT(s)" + " " + str(port[0]) + " " + port[1]     # Sets the output and converts int to char for port number
			
			if len(scanned_hosts.get_open_ports()) != 0: #identifies open ports
				
				print_output = "IP" + " " + output + " SOURCE_NET " + sourceNetwork + " DEST_NET" + " " + destinations[dest] + " " + dest + "\n"   # This defines the print output format
				
				print print_output    # Printed output to std out
				
				outPutFile.write(print_output)    # Printed output to file with source network name and source network ip

				if allReadyPrinted == False:
					document.add_heading('Network Segmentation Validation Results', 0)
					#document.add_heading('Segmentation Results', level=1)
					document.add_paragraph('The following are the segmentation and routing validation results identified from the source network' + ' ' + sourceNetworkName + ' ' + sourceNetwork + '.')
					allReadyPrinted = True
				
				document.add_paragraph(print_output, style="List Number")

				
else:
	print nm.stderr    # Condition for script error



outPutFile.close() # Closing the output file
document.save(sourceNetworkName + '_' + sourceNetwork.replace('/','-') + '.docx')










